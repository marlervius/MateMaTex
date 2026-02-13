"""
LaTeX → Word (DOCX) conversion.

Uses python-docx for programmatic construction.
Strips LaTeX commands and renders math as plain text for now;
a future version can use OMML for proper equation rendering.
"""

from __future__ import annotations

import io
import re

import structlog

logger = structlog.get_logger()


def _strip_latex_commands(text: str) -> str:
    """Convert LaTeX to readable plain text for Word export."""
    # Remove comments
    text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)

    # Remove preamble
    doc_begin = text.find(r'\begin{document}')
    if doc_begin >= 0:
        text = text[doc_begin + len(r'\begin{document}'):]
    doc_end = text.find(r'\end{document}')
    if doc_end >= 0:
        text = text[:doc_end]

    # Remove environments (keep content)
    text = re.sub(r'\\begin\{[^}]*\}(?:\{[^}]*\})?', '', text)
    text = re.sub(r'\\end\{[^}]*\}', '', text)

    # Replace common commands
    replacements = {
        r'\maketitle': '',
        r'\newpage': '',
        r'\noindent': '',
        r'\vspace{': '',
        r'\hspace{': '',
        r'\\': '\n',
        r'\item': '• ',
        r'\textbf{': '',
        r'\textit{': '',
        r'\emph{': '',
        r'\sffamily': '',
        r'\bfseries': '',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # Extract section titles
    text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n\n\1\n', text)
    text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n\1\n', text)
    text = re.sub(r'\\title\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\author\{([^}]*)\}', r'\1', text)

    # Simplify math
    text = re.sub(r'\$\$([^$]*)\$\$', r' \1 ', text)
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    text = re.sub(r'\\\[([^]]*)\\\]', r' \1 ', text)
    text = re.sub(r'\\\(([^)]*)\\\)', r'\1', text)
    text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', text)
    text = re.sub(r'\\cdot', r'·', text)
    text = re.sub(r'\\times', r'×', text)
    text = re.sub(r'\\pm', r'±', text)
    text = re.sub(r'\\leq', r'≤', text)
    text = re.sub(r'\\geq', r'≥', text)
    text = re.sub(r'\\neq', r'≠', text)
    text = re.sub(r'\\pi', r'π', text)
    text = re.sub(r'\\alpha', r'α', text)
    text = re.sub(r'\\beta', r'β', text)

    # Clean remaining commands
    text = re.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)
    text = re.sub(r'[{}]', '', text)

    # Clean whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def latex_to_docx(
    latex_content: str,
    title: str = "Oppgaveark",
    include_solutions: bool = True,
) -> bytes:
    """
    Convert LaTeX content to a Word document.

    Returns the DOCX file as bytes.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generert av MateMaTeX AI")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Spacer

    # Convert content
    plain_text = _strip_latex_commands(latex_content)

    # Split into sections
    sections = re.split(r'\n\n+', plain_text)

    for section in sections:
        section = section.strip()
        if not section:
            continue

        # Check if it's a heading-like line (short, no bullet)
        lines = section.split('\n')
        first_line = lines[0].strip()

        if len(lines) == 1 and len(first_line) < 80 and not first_line.startswith('•'):
            # Likely a heading
            if any(word in first_line.lower() for word in ['oppgave', 'eksempel', 'definisjon', 'løsning']):
                doc.add_heading(first_line, level=2)
                continue

        # Regular paragraph
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('• '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
